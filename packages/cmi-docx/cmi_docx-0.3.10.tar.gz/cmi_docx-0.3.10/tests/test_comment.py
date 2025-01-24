"""Tests for the comment module."""

import docx
from docx.opc import constants as docx_constants

from cmi_docx import comment


def test_add_comment_single() -> None:
    """Tests adding a comment to a single entity."""
    document = docx.Document()
    para = document.add_paragraph("This is a sample paragraph.")
    author = "Grievous"
    message = "Ah, General Kenobi."

    comment.add_comment(document, para, author, message)
    para_comment = para.part.part_related_by(docx_constants.RELATIONSHIP_TYPE.COMMENTS)

    assert author in para_comment.blob.decode()
    assert message in para_comment.blob.decode()


def test_add_comment_range() -> None:
    """Tests adding a comment with a range of entities."""
    document = docx.Document()
    para = document.add_paragraph("This is a sample paragraph.")
    run_start = para.add_run(" run 3!")
    run_end = para.add_run(" run 4!")
    author = "Grievous"
    message = "Ah, General Kenobi."

    comment.add_comment(document, (run_start, run_end), author, message)
    comment_start = run_start.part.part_related_by(
        docx_constants.RELATIONSHIP_TYPE.COMMENTS
    )
    comment_end = run_end.part.part_related_by(
        docx_constants.RELATIONSHIP_TYPE.COMMENTS
    )

    assert author in comment_start.blob.decode()
    assert message in comment_start.blob.decode()
    assert author in comment_end.blob.decode()
    assert message in comment_end.blob.decode()
